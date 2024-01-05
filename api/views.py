from django.shortcuts import redirect
from rest_framework.generics import CreateAPIView, ListAPIView
from rest_framework.response import Response
from rest_framework.views import APIView
from docx import Document  # third

from .models import ConvertTextModel, ConvertFileModel
from .serializers import ConvertTextSerializer, ConvertFileSerializer
from .convert import to_latin, to_ciril


# Create your views here.
class HomeViews(ListAPIView):
    queryset = ConvertTextModel.objects.order_by('-date')[:1]
    serializer_class = ConvertTextSerializer


class AddTextView(CreateAPIView):
    queryset = ConvertTextModel.objects.all()
    serializer_class = ConvertTextSerializer

    def create(self, request, *args, **kwargs):
        response = super().create(request, *args, **kwargs)
        return redirect('converter')


class AddFileView(CreateAPIView):
    queryset = ConvertFileModel.objects.all()
    serializer_class = ConvertFileSerializer

    def create(self, request, *args, **kwargs):
        response = super().create(request, *args, **kwargs)
        return redirect('open-file')


class LatinOrCyrillicView(APIView):
    def get(self, request, *args, **kwargs):
        last_text = ConvertTextModel.objects.last().text
        to = ConvertTextModel.objects.last().to
        try:
            if to == "to_Latin":
                convert = to_latin(last_text)
                return Response({'Original Text': last_text, 'Cyrillic to Latin': convert})
            elif to == 'to_Cyrillic':
                convert = to_ciril(last_text)
                return Response({'Original Text': last_text, 'Latin to Cyrillic': convert})
        except ImportError:  # KeyError
            return Response({'result': "Siz konvert qilishda adashdingiz!"})


class OpenFileView(APIView):
    def get(self, request, *args, **kwargs):
        last_text = ConvertFileModel.objects.last().file.name
        to = ConvertFileModel.objects.last().to

        if last_text[-4:] == 'docx':
            # print('Docx ga kirildi!')
            fayl = last_text
            document = Document(fayl)
            res = """"""
            for paragraph in document.paragraphs:
                res += paragraph.text
            try:
                if to == "to_Latin":
                    convert = to_latin(res)

                    for i in range(len(document.paragraphs)):
                        document.paragraphs[i].clear()
                    document.add_paragraph(convert)
                    document.save(last_text)
                    return Response({'Original Text': f"https://alphabet.pythonanywhere.com/{last_text}", 'Cyrillic to Latin': convert})

                elif to == 'to_Cyrillic':
                    convert = to_ciril(res)
                    # print(convert)

                    for i in range(len(document.paragraphs)):
                        document.paragraphs[i].clear()

                    document.add_paragraph(convert)
                    document.save(last_text)
                    return Response({'Original Text': f"https://alphabet.pythonanywhere.com/{last_text}", 'Latin to Cyrillic': convert})

            except KeyError:  # KeyError
                return Response({'result': "Siz konvert qilishda adashdingiz! Yoki Alfabetda mavjud bo'lmagan belgi "
                                           "kiritilgan!"})

        else:
            print('Start!')
            with open(last_text, 'r+', encoding='utf-8') as fayl:
                text = fayl.readlines()
                # print(text)
                for i in text:
                    res = i.strip()
                    # print(res)
            try:
                print('Kirildi!')
                if to == "to_Latin":
                    convert = to_latin(res)
                    # print(convert)
                    print(last_text)

                    with open(last_text, 'w+', encoding='utf-8') as fayl:
                        for convert in fayl:
                            fayl.writelines(convert)
                            print(fayl)

                    return Response({'Original Text':  f"https://alphabet.pythonanywhere.com/{last_text}", 'Cyrillic to Latin': convert})

                elif to == 'to_Cyrillic':
                    convert = to_ciril(res)
                    with open(last_text, 'w+', encoding='utf-8') as fayl:
                        for convert in fayl:
                            fayl.writelines(convert)
                            print(fayl)
                    return Response({'Original Text':  f"https://alphabet.pythonanywhere.com/{last_text}", 'Latin to Cyrillic': convert})
            except KeyError:  # KeyError
                return Response({'result': "Siz konvert qilishda adashdingiz!"})

        return Response(last_text)
